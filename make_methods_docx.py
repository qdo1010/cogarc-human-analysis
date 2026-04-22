"""Convert a plain-text findings / methods file into a .docx.

Preserves headings (numbered sections), bullet-like sub-items, and
monospaced blocks (indented code/paths / tables).

CLI:
    python make_methods_docx.py \
        [--src docs/motor_vs_cognitive_methods.txt] \
        [--dst docs/motor_vs_cognitive_methods.docx] \
        [--title "Motor slip vs cognitive error"] \
        [--subtitle "Classification of edit-level errors ..."]
"""

from __future__ import annotations

import _paths  # noqa: F401  (sys.path bootstrap, safe under both invocations)

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

DEFAULT_SRC = Path("docs/motor_vs_cognitive_methods.txt")
DEFAULT_DST = Path("docs/motor_vs_cognitive_methods.docx")
DEFAULT_TITLE = "Motor slip vs cognitive error"
DEFAULT_SUBTITLE = "Classification of edit-level errors in the CogARC Experiment 2 dataset"


# Section headings in the source are of the form:
#   "1. DATA" / "2. PER-EDIT SIGNALS" ... followed by a dashed underline.
SECTION_RE = re.compile(r"^\s*(\d+)\.\s+(.+?)\s*$")
SUBITEM_RE = re.compile(r"^\s{4}([a-z])\)\s+(.*)$")      # "    a) ..."
ROMAN_RE = re.compile(r"^\s{4}(i{1,3}|iv|v)\)\s+(.*)$")   # "    i) ..."


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", default=str(DEFAULT_SRC),
                    help="Path to the .txt source file.")
    ap.add_argument("--dst", default=None,
                    help="Output .docx path. Default: same stem as --src.")
    ap.add_argument("--title", default=DEFAULT_TITLE)
    ap.add_argument("--subtitle", default=DEFAULT_SUBTITLE)
    args = ap.parse_args()

    src_path = Path(args.src)
    dst_path = Path(args.dst) if args.dst else src_path.with_suffix(".docx")

    text = src_path.read_text()
    lines = text.splitlines()

    doc = Document()

    title = doc.add_heading(args.title, level=0)
    for run in title.runs:
        run.font.size = Pt(22)

    subtitle = doc.add_paragraph(args.subtitle)
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(13)

    # Tighter defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Simple state machine over the source lines
    i = 0
    skip_until_non_dash = False
    in_code_block = False

    def flush_code(buffer):
        if not buffer:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(18)
        for k, line in enumerate(buffer):
            if k > 0:
                p.add_run().add_break()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

    code_buffer: list[str] = []

    # Skip the leading banner (first two non-empty title lines + dashed rule)
    # which we've already rendered via the Title/Subtitle above.
    skip_leading = True

    def _is_sep(s: str) -> bool:
        s = s.strip()
        return bool(s) and set(s) <= {"-", "="} and len(s) >= 3

    while i < len(lines):
        ln = lines[i]

        if skip_leading:
            if ln.startswith("====="):
                skip_leading = False
            i += 1
            continue

        # A dashed separator line on its own is a decoration — skip.
        if _is_sep(ln):
            i += 1
            continue

        # A flush-left line followed by a dashed underline is a heading.
        # (Requiring flush-left avoids misclassifying indented table
        # separators like "    ----" that follow a column header.)
        # Numbered "1. SECTION" lines become Heading 1; all other heading
        # candidates become Heading 2.
        if (ln.strip() and not ln.startswith(" ")
                and i + 1 < len(lines) and _is_sep(lines[i + 1])):
            flush_code(code_buffer); code_buffer = []
            in_code_block = False
            m = SECTION_RE.match(ln)
            if m:
                heading_text = f"{m.group(1)}. {m.group(2).title()}"
                level = 1
            else:
                heading_text = ln.strip()
                level = 2
            doc.add_heading(heading_text, level=level)
            i += 2
            continue

        # Blank line: end any current code block (we treat a blank line as a
        # paragraph break, so we accumulate code only across adjacent indented
        # lines).
        if not ln.strip():
            flush_code(code_buffer); code_buffer = []
            in_code_block = False
            i += 1
            continue

        # Indented lines (4+ leading spaces) are typically code/path/constants.
        # Exception: sub-items like "    a) ..." or "    i) ..." which should
        # render as regular paragraphs with a small indent.
        sub = SUBITEM_RE.match(ln) or ROMAN_RE.match(ln)
        if sub:
            flush_code(code_buffer); code_buffer = []
            in_code_block = False
            letter, body = sub.group(1), sub.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(f"{letter})  ")
            r.bold = True
            p.add_run(body)
            # Consume subsequent continuation lines (indented more)
            j = i + 1
            while j < len(lines) and lines[j].startswith("       ") and lines[j].strip():
                p.add_run(" " + lines[j].strip())
                j += 1
            i = j
            continue

        # Hyphen bullet list at any indent ("    - foo"). CHECK BEFORE
        # the generic indented-code branch so bulleted lines don't get
        # misclassified as code.
        if ln.lstrip().startswith("- "):
            flush_code(code_buffer); code_buffer = []
            p = doc.add_paragraph(ln.lstrip()[2:], style="List Bullet")
            p.paragraph_format.left_indent = Pt(18)
            # Absorb continuation lines that are more deeply indented
            j = i + 1
            while j < len(lines) and lines[j].startswith("       ") and lines[j].strip() \
                    and not lines[j].lstrip().startswith("- "):
                p.add_run(" " + lines[j].strip())
                j += 1
            i = j
            continue

        if ln.startswith("    "):
            code_buffer.append(ln[4:] if ln.startswith("    ") else ln)
            in_code_block = True
            i += 1
            continue

        # Normal paragraph — coalesce wrapped lines until blank or a new
        # structural element.
        flush_code(code_buffer); code_buffer = []
        buf = [ln.rstrip()]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if not nxt.strip():
                break
            if SECTION_RE.match(nxt) and j + 1 < len(lines) and set(lines[j + 1].strip()) == {"-"}:
                break
            if SUBITEM_RE.match(nxt) or ROMAN_RE.match(nxt):
                break
            if nxt.startswith("    "):
                break
            if nxt.lstrip().startswith("- "):
                break
            buf.append(nxt.rstrip())
            j += 1
        doc.add_paragraph(" ".join(s.strip() for s in buf))
        i = j

    flush_code(code_buffer)

    dst_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst_path)
    print(f"wrote {dst_path}")


if __name__ == "__main__":
    main()
